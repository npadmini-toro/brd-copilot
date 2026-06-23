import json
import io
import streamlit as st
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

# ----- Config -----
st.set_page_config(page_title="BRD Copilot", page_icon="📋", layout="wide")

MODELS = {
    "Sonnet (balanced, default)": "claude-sonnet-4-6",
    "Opus (deepest analysis)": "claude-opus-4-8",
    "Haiku (fast/cheap)": "claude-haiku-4-5-20251001",
}

# ----- Secrets check -----
missing = [k for k in ["ANTHROPIC_API_KEY"] if k not in st.secrets]
if missing:
    st.error(f"Missing secret(s): {', '.join(missing)}. Add them in .streamlit/secrets.toml (local) or the Streamlit Cloud dashboard.")
    st.stop()

client = Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])

# ----- The BRD contract -----
SYSTEM_PROMPT = """You are a senior Business Analyst. Given a business problem, produce a complete, traceable Business Requirements Document (BRD).

Return ONLY valid JSON (no prose, no markdown fences) matching this schema:
{
  "scope_ok": true,
  "redirect_message": "",
  "document_control": {"title": "", "author": "", "version": "1.0", "date": "", "status": "Draft"},
  "business_context": {"problem_statement": "", "objectives": [""], "success_metrics": [""]},
  "scope": {"in_scope": [""], "out_of_scope": [""], "assumptions": [""], "constraints": [""]},
  "stakeholders": [{"role": "", "interest": "", "raci": ""}],
  "functional_requirements": [{"id": "FR-01", "statement": "", "priority": "Must|Should|Could|Won't", "rationale": ""}],
  "non_functional_requirements": [{"id": "NFR-01", "category": "", "statement": ""}],
  "user_stories": [{"id": "US-01", "story": "As a ..., I want ..., so that ...", "acceptance_criteria": ["Given ..., when ..., then ..."], "linked_fr": "FR-01"}],
  "process_flow": {"current_state": [""], "future_state": [""]},
  "traceability_matrix": [{"objective": "", "fr_id": "FR-01", "us_id": "US-01", "tc_id": "TC-01"}],
  "test_scenarios": [{"id": "TC-01", "scenario": "", "steps": [""], "expected_result": "", "linked_requirement": "FR-01"}],
  "uat": {"entry_criteria": [""], "exit_criteria": [""]},
  "gap_log": [{"item": "", "issue_type": "Vague|Incomplete|Untestable|Conflicting|Missing", "clarifying_question": ""}],
  "data_validation_sql": [{"requirement": "", "sql": "", "proves": ""}]
}

Rules:
- If the input is NOT a software/business requirements problem, set scope_ok=false, fill redirect_message, leave other fields empty.
- Use MoSCoW for functional requirement priority.
- Every functional requirement must trace to at least one user story and one test case via shared IDs.
- Acceptance criteria must be in Given/When/Then form and testable.
- In gap_log, surface anything vague, incomplete, untestable, or contradictory in the source request, with a specific clarifying question for each.
- Only include data_validation_sql for requirements that involve data rules (eligibility, thresholds, counts). Use standard ANSI SQL.
- In data_validation_sql, never use double-quote characters. Use unquoted identifiers and single quotes for string literals only, so the surrounding JSON stays valid.
"""

def extract_json(text):
    s, e = text.find("{"), text.rfind("}")
    return text[s:e + 1]

# ----- Word export -----
def _add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    r = run._r
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-1" \\h \\z \\u'   # only Heading 1 entries
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = "Right-click here and choose 'Update Field' to build the table of contents."
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    for el in (begin, instr, sep, placeholder, end):
        r.append(el)

def build_docx(brd):
    doc = Document()
    dc = brd.get("document_control", {})

    # --- Title page ---
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(dc.get("title", "Business Requirements Document"))
    run.bold = True; run.font.size = Pt(28)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Business Requirements Document").font.size = Pt(16)
    for _ in range(2):
        doc.add_paragraph()
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Version {dc.get('version','1.0')}     Status: {dc.get('status','Draft')}\n"
        f"Author: {dc.get('author','')}\n"
        f"Date: {dc.get('date','')}"
    ).font.size = Pt(12)
    doc.add_page_break()

    # --- TOC page ---
    doc.add_heading("Table of Contents", level=1)
    _add_toc(doc)
    doc.add_page_break()

    def heading(t): doc.add_heading(t, level=1)
    def bullets(items):
        for x in items:
            doc.add_paragraph(str(x), style="List Bullet")

    bc = brd.get("business_context", {})
    heading("1. Business Context")
    doc.add_paragraph(bc.get("problem_statement", ""))
    doc.add_paragraph("Objectives:"); bullets(bc.get("objectives", []))
    doc.add_paragraph("Success Metrics:"); bullets(bc.get("success_metrics", []))

    sc = brd.get("scope", {})
    heading("2. Scope")
    doc.add_paragraph("In Scope:"); bullets(sc.get("in_scope", []))
    doc.add_paragraph("Out of Scope:"); bullets(sc.get("out_of_scope", []))
    doc.add_paragraph("Assumptions:"); bullets(sc.get("assumptions", []))
    doc.add_paragraph("Constraints:"); bullets(sc.get("constraints", []))

    def table(rows, cols):
        if not rows: return
        tbl = doc.add_table(rows=1, cols=len(cols)); tbl.style = "Light Grid Accent 1"
        for i, c in enumerate(cols): tbl.rows[0].cells[i].text = c
        for row in rows:
            cells = tbl.add_row().cells
            for i, c in enumerate(cols): cells[i].text = str(row.get(c, ""))

    heading("3. Stakeholders")
    table(brd.get("stakeholders", []), ["role", "interest", "raci"])
    heading("4. Functional Requirements")
    table(brd.get("functional_requirements", []), ["id", "statement", "priority", "rationale"])
    heading("5. Non-Functional Requirements")
    table(brd.get("non_functional_requirements", []), ["id", "category", "statement"])

    heading("6. User Stories")
    for us in brd.get("user_stories", []):
        doc.add_paragraph(f"{us.get('id','')} (links to {us.get('linked_fr','')})", style="Heading 3")
        doc.add_paragraph(us.get("story", ""))
        for ac in us.get("acceptance_criteria", []):
            doc.add_paragraph(ac, style="List Bullet")

    heading("7. Requirements Traceability Matrix")
    table(brd.get("traceability_matrix", []), ["objective", "fr_id", "us_id", "tc_id"])

    heading("8. Test Scenarios & UAT")
    for tc in brd.get("test_scenarios", []):
        doc.add_paragraph(f"{tc.get('id','')} — {tc.get('scenario','')} (validates {tc.get('linked_requirement','')})", style="Heading 3")
        for s in tc.get("steps", []):
            doc.add_paragraph(s, style="List Number")
        doc.add_paragraph(f"Expected: {tc.get('expected_result','')}")
    uat = brd.get("uat", {})
    doc.add_paragraph("UAT Entry Criteria:"); bullets(uat.get("entry_criteria", []))
    doc.add_paragraph("UAT Exit Criteria:"); bullets(uat.get("exit_criteria", []))

    heading("9. Gap & Ambiguity Log")
    table(brd.get("gap_log", []), ["item", "issue_type", "clarifying_question"])

    heading("10. Data Validation SQL")
    for q in brd.get("data_validation_sql", []):
        doc.add_paragraph(q.get("requirement", ""), style="Heading 3")
        doc.add_paragraph(q.get("sql", ""))
        doc.add_paragraph(f"Proves: {q.get('proves','')}")

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ----- Excel export -----
def build_xlsx(brd):
    wb = Workbook(); wb.remove(wb.active)
    def sheet(name, rows, cols):
        ws = wb.create_sheet(name[:31]); ws.append(cols)
        for r in rows:
            ws.append([str(r.get(c, "")) for c in cols])
    sheet("Functional Reqs", brd.get("functional_requirements", []), ["id", "statement", "priority", "rationale"])
    sheet("Non-Functional Reqs", brd.get("non_functional_requirements", []), ["id", "category", "statement"])
    sheet("Traceability", brd.get("traceability_matrix", []), ["objective", "fr_id", "us_id", "tc_id"])
    sheet("Test Scenarios", brd.get("test_scenarios", []), ["id", "scenario", "expected_result", "linked_requirement"])
    sheet("Gap Log", brd.get("gap_log", []), ["item", "issue_type", "clarifying_question"])
    if not wb.sheetnames:
        wb.create_sheet("BRD")
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()

# ----- UI -----
st.title("📋 BRD Copilot")
st.caption("Turn a business problem into a complete, traceable Business Requirements Document.")

with st.sidebar:
    model_label = st.selectbox("Model", list(MODELS.keys()))
    model = MODELS[model_label]

problem = st.text_area(
    "Describe the business problem or feature request",
    height=160,
    placeholder="e.g. Members can't see the status of their submitted claims and call support constantly. We want a self-service claims status portal.",
)

go = st.button("Generate BRD", type="primary")

# ----- Generate -----
if go:
    if not problem.strip():
        st.warning("Please describe a business problem first.")
        st.stop()
    with st.spinner("Analyzing requirements and drafting the BRD..."):
        resp = client.messages.create(
            model=model, max_tokens=16000, system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": problem}],
        )
        candidate = extract_json(resp.content[0].text)
        try:
            brd = json.loads(candidate)
        except json.JSONDecodeError:
            fix = client.messages.create(
                model=model, max_tokens=16000,
                system="You fix invalid JSON. Return ONLY corrected, valid JSON — no prose, no code fences. Escape or remove stray quotes inside string values.",
                messages=[{"role": "user", "content": candidate}],
            )
            try:
                brd = json.loads(extract_json(fix.content[0].text))
            except json.JSONDecodeError as e:
                st.error(f"Could not parse JSON even after repair: {e}")
                st.session_state.pop("brd", None); st.stop()
                brd.setdefault("document_control", {})["date"] = date.today().strftime("%d %B %Y")
        st.session_state["brd"] = brd
       

# ----- Render -----
if "brd" in st.session_state:
    brd = st.session_state["brd"]
    if not brd.get("scope_ok", False):
        st.warning(brd.get("redirect_message", "That input doesn't look like a requirements problem."))
        st.stop()

    dc = brd.get("document_control", {})
    st.success(f"BRD generated: {dc.get('title', 'Untitled')}")

    c1, c2 = st.columns(2)
    c1.download_button("⬇️ Download Word (.docx)", build_docx(brd),
        file_name="BRD.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    c2.download_button("⬇️ Download Excel (.xlsx)", build_xlsx(brd),
        file_name="BRD_traceability.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    tabs = st.tabs([
        "📄 Overview", "📋 Requirements", "👤 User Stories",
        "🔀 Process Flow", "🔗 Traceability", "✅ Tests & UAT",
        "⚠️ Gaps", "🗄️ Validation SQL",
    ])

    with tabs[0]:
        c1, c2, c3 = st.columns(3)
        c1.metric("Version", dc.get("version", "—"))
        c2.metric("Status", dc.get("status", "—"))
        c3.metric("Date", dc.get("date", "—"))
        bc = brd.get("business_context", {})
        st.subheader("Problem Statement"); st.write(bc.get("problem_statement", ""))
        st.subheader("Objectives")
        for o in bc.get("objectives", []): st.markdown(f"- {o}")
        st.subheader("Success Metrics")
        for m in bc.get("success_metrics", []): st.markdown(f"- {m}")
        sc = brd.get("scope", {})
        c1, c2 = st.columns(2)
        with c1:
            st.subheader("In Scope")
            for x in sc.get("in_scope", []): st.markdown(f"- {x}")
            st.subheader("Assumptions")
            for x in sc.get("assumptions", []): st.markdown(f"- {x}")
        with c2:
            st.subheader("Out of Scope")
            for x in sc.get("out_of_scope", []): st.markdown(f"- {x}")
            st.subheader("Constraints")
            for x in sc.get("constraints", []): st.markdown(f"- {x}")
        st.subheader("Stakeholders")
        st.dataframe(brd.get("stakeholders", []), use_container_width=True, hide_index=True)

    with tabs[1]:
        st.subheader("Functional Requirements")
        st.dataframe(brd.get("functional_requirements", []), use_container_width=True, hide_index=True)
        st.subheader("Non-Functional Requirements")
        st.dataframe(brd.get("non_functional_requirements", []), use_container_width=True, hide_index=True)

    with tabs[2]:
        for us in brd.get("user_stories", []):
            with st.expander(f"{us.get('id', '')} — links to {us.get('linked_fr', '')}"):
                st.write(us.get("story", ""))
                st.markdown("**Acceptance Criteria**")
                for ac in us.get("acceptance_criteria", []): st.markdown(f"- {ac}")

    with tabs[3]:
        pf = brd.get("process_flow", {})
        c1, c2 = st.columns(2)
        with c1:
            st.subheader("Current State")
            for i, step in enumerate(pf.get("current_state", []), 1): st.markdown(f"{i}. {step}")
        with c2:
            st.subheader("Future State")
            for i, step in enumerate(pf.get("future_state", []), 1): st.markdown(f"{i}. {step}")

    with tabs[4]:
        st.subheader("Requirements Traceability Matrix")
        st.caption("Objective → Functional Requirement → User Story → Test Case")
        st.dataframe(brd.get("traceability_matrix", []), use_container_width=True, hide_index=True)

    with tabs[5]:
        st.subheader("Test Scenarios")
        for tc in brd.get("test_scenarios", []):
            with st.expander(f"{tc.get('id', '')} — {tc.get('scenario', '')} (validates {tc.get('linked_requirement', '')})"):
                st.markdown("**Steps**")
                for i, s in enumerate(tc.get("steps", []), 1): st.markdown(f"{i}. {s}")
                st.markdown(f"**Expected Result:** {tc.get('expected_result', '')}")
        uat = brd.get("uat", {})
        c1, c2 = st.columns(2)
        with c1:
            st.subheader("UAT Entry Criteria")
            for x in uat.get("entry_criteria", []): st.markdown(f"- {x}")
        with c2:
            st.subheader("UAT Exit Criteria")
            for x in uat.get("exit_criteria", []): st.markdown(f"- {x}")

    with tabs[6]:
        st.subheader("Gap & Ambiguity Log")
        st.caption("Vague, incomplete, untestable, or conflicting items found in the source request.")
        st.dataframe(brd.get("gap_log", []), use_container_width=True, hide_index=True)

    with tabs[7]:
        st.subheader("Data Validation SQL")
        for q in brd.get("data_validation_sql", []):
            st.markdown(f"**{q.get('requirement', '')}**")
            st.code(q.get("sql", ""), language="sql")
            st.caption(f"Proves: {q.get('proves', '')}")
            st.divider()